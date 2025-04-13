import sys
import os
import pickle
import numpy as np
import pandas as pd
import argparse
import joblib
import re
from math import log2
import hashlib
from sklearn.feature_selection import SelectKBest, mutual_info_classif

class UnifiedMalwarePredictor:
    """Unified class to predict malware across multiple file types, activating only the relevant function"""

    def __init__(self, pdf_model_path="models/rf_model.pkl",
                 exe_model_path="models/exe_model.pkl",
                 multi_model_path="models/malware_models.pkl",
                 small_file_threshold=10240,
                 zero_features_threshold=0.7,
                 top_k_features=20):
        self.pdf_model = None
        self.exe_model = None
        self.multi_models = None
        self.multi_feature_names = None
        self.selected_features = None
        self.small_file_threshold = small_file_threshold
        self.zero_features_threshold = zero_features_threshold
        self.top_k_features = top_k_features
        self.pdf_model_path = pdf_model_path
        self.exe_model_path = exe_model_path
        self.multi_model_path = multi_model_path
        self.magic = None
        self.fitz = None
        self.pdfplumber = None
        self.pefile = None
        self.olefile = None
        self.docx = None

    def load_pdf_model(self):
        """Load the PDF model"""
        if self.pdf_model is None:
            try:
                if os.path.exists(self.pdf_model_path):
                    self.pdf_model = joblib.load(self.pdf_model_path)
                    print(f"✅ Loaded PDF model from {self.pdf_model_path}")
                else:
                    print(f"⚠️ PDF model not found at {self.pdf_model_path}")
            except Exception as e:
                print(f"❌ Error loading PDF model: {str(e)}")
                self.pdf_model = None

    def load_exe_model(self):
        """Load the EXE model"""
        if self.exe_model is None:
            try:
                if os.path.exists(self.exe_model_path):
                    self.exe_model = joblib.load(self.exe_model_path)
                    print(f"✅ Loaded EXE model from {self.exe_model_path}")
                else:
                    print(f"⚠️ EXE model not found at {self.exe_model_path}")
            except Exception as e:
                print(f"❌ Error loading EXE model: {str(e)}")
                self.exe_model = None

    def load_multi_model(self):
        """Load the multi-format model"""
        if self.multi_models is None:
            try:
                if os.path.exists(self.multi_model_path):
                    with open(self.multi_model_path, "rb") as f:
                        self.multi_models = pickle.load(f)
                    print(f"✅ Loaded multi-format models from {self.multi_model_path}")
                    self.multi_feature_names = [
                        'e_magic', 'e_cblp', 'e_cp', 'e_crlc', 'e_cparhdr', 'e_minalloc', 'e_maxalloc',
                        'e_ss', 'e_sp', 'e_csum', 'e_ip', 'e_cs', 'e_lfarlc', 'e_ovno', 'e_oemid',
                        'e_oeminfo', 'e_lfanew', 'Machine', 'NumberOfSections', 'TimeDateStamp',
                        'PointerToSymbolTable', 'NumberOfSymbols', 'SizeOfOptionalHeader', 'Characteristics',
                        'Magic', 'MajorLinkerVersion', 'MinorLinkerVersion', 'SizeOfCode',
                        'SizeOfInitializedData', 'SizeOfUninitializedData', 'AddressOfEntryPoint',
                        'BaseOfCode', 'ImageBase', 'SectionAlignment', 'FileAlignment',
                        'MajorOperatingSystemVersion', 'MinorOperatingSystemVersion', 'MajorImageVersion',
                        'MinorImageVersion', 'MajorSubsystemVersion', 'MinorSubsystemVersion', 'Reserved1',
                        'SizeOfImage', 'SizeOfHeaders', 'CheckSum', 'Subsystem', 'DllCharacteristics',
                        'SizeOfStackReserve', 'SizeOfHeapReserve', 'SizeOfHeapCommit', 'LoaderFlags',
                        'NumberOfRvaAndSizes', 'file_size', 'entropy', 'contains_macros', 'suspicious_streams',
                        'object_count', 'javascript_count', 'urls_count', 'suspicious_strings',
                        'encryption_indicator', 'obfuscation_level', 'contains_executable', 'unusual_metadata'
                    ]
                else:
                    print(f"⚠️ Multi-format model not found at {self.multi_model_path}")
            except Exception as e:
                print(f"❌ Error loading multi-format model: {str(e)}")
                self.multi_models = None

    def detect_javascript(self, text):
        """Detect JavaScript presence in text"""
        js_keywords = ['app.alert', 'this.submitForm', 'this.getField', '/JavaScript']
        return any(kw in text for kw in js_keywords)

    def extract_pdf_features(self, file_path):
        """Extract features from a PDF file"""
        try:
            if self.fitz is None:
                import fitz
                self.fitz = fitz
            if self.pdfplumber is None:
                import pdfplumber
                self.pdfplumber = pdfplumber

            doc = self.fitz.open(file_path)
            file_size = os.path.getsize(file_path)
            num_pages = len(doc)
            has_openaction = 0
            has_launch_action = 0
            has_embedded_file = 0
            has_uri = 0
            has_js = 0
            num_objects = 0
            num_streams = 0
            encoded_stream_ratio = 0.0
            entropy = 0.0

            try:
                catalog = doc.pdf_catalog()
                has_openaction = int("/OpenAction" in catalog if isinstance(catalog, dict) else False)
            except:
                has_openaction = 0

            all_text = ""
            for page in doc:
                text = page.get_text()
                all_text += text
                links = page.get_links()
                if links:
                    for link in links:
                        if "uri" in link:
                            has_uri = 1
                if '/Launch' in page.get_text("text"):
                    has_launch_action = 1

            has_js = int(self.detect_javascript(all_text))
            if doc.embfile_count() > 0:
                has_embedded_file = 1

            with open(file_path, 'rb') as f:
                data = f.read()
                prob = [float(data.count(byte)) / len(data) for byte in set(data)]
                entropy = -sum([p * log2(p) for p in prob if p > 0])

            with self.pdfplumber.open(file_path) as pdf:
                text = "".join([p.extract_text() or "" for p in pdf.pages])
                num_streams = text.count("stream")
                num_objects = text.count("obj")
            encoded_stream_ratio = num_streams / num_objects if num_objects > 0 else 0.0

            features = {
                'has_js': has_js,
                'has_openaction': has_openaction,
                'has_embedded_file': has_embedded_file,
                'has_uri': has_uri,
                'has_launch_action': has_launch_action,
                'entropy': round(entropy, 4),
                'num_objects': num_objects,
                'num_streams': num_streams,
                'num_pages': num_pages,
                'encoded_stream_ratio': round(encoded_stream_ratio, 4),
                'file_size': file_size
            }
            return pd.DataFrame([features])

        except ImportError as e:
            print(f"❌ Missing required PDF libraries: {str(e)}")
            return None
        except Exception as e:
            print(f"❌ Error extracting PDF features: {e}")
            return None

    def extract_exe_features(self, file_path):
        """Extract features from an EXE/DLL file"""
        try:
            if self.pefile is None:
                import pefile
                self.pefile = pefile

            pe = self.pefile.PE(file_path)
            features = {
                'e_magic': pe.DOS_HEADER.e_magic,
                'e_cblp': pe.DOS_HEADER.e_cblp,
                'e_cp': pe.DOS_HEADER.e_cp,
                'e_crlc': pe.DOS_HEADER.e_crlc,
                'e_cparhdr': pe.DOS_HEADER.e_cparhdr,
                'e_minalloc': pe.DOS_HEADER.e_minalloc,
                'e_maxalloc': pe.DOS_HEADER.e_maxalloc,
                'e_ss': pe.DOS_HEADER.e_ss,
                'e_sp': pe.DOS_HEADER.e_sp,
                'e_csum': pe.DOS_HEADER.e_csum,
                'e_ip': pe.DOS_HEADER.e_ip,
                'e_cs': pe.DOS_HEADER.e_cs,
                'e_lfarlc': pe.DOS_HEADER.e_lfarlc,
                'e_ovno': pe.DOS_HEADER.e_ovno,
                'e_oemid': pe.DOS_HEADER.e_oemid,
                'e_oeminfo': pe.DOS_HEADER.e_oeminfo,
                'e_lfanew': pe.DOS_HEADER.e_lfanew,
                'Machine': pe.FILE_HEADER.Machine,
                'NumberOfSections': pe.FILE_HEADER.NumberOfSections,
                'TimeDateStamp': pe.FILE_HEADER.TimeDateStamp,
                'PointerToSymbolTable': pe.FILE_HEADER.PointerToSymbolTable,
                'NumberOfSymbols': pe.FILE_HEADER.NumberOfSymbols,
                'SizeOfOptionalHeader': pe.FILE_HEADER.SizeOfOptionalHeader,
                'Characteristics': pe.FILE_HEADER.Characteristics,
                'Magic': pe.OPTIONAL_HEADER.Magic,
                'MajorLinkerVersion': pe.OPTIONAL_HEADER.MajorLinkerVersion,
                'MinorLinkerVersion': pe.OPTIONAL_HEADER.MinorLinkerVersion,
                'SizeOfCode': pe.OPTIONAL_HEADER.SizeOfCode,
                'SizeOfInitializedData': pe.OPTIONAL_HEADER.SizeOfInitializedData,
                'SizeOfUninitializedData': pe.OPTIONAL_HEADER.SizeOfUninitializedData,
                'AddressOfEntryPoint': pe.OPTIONAL_HEADER.AddressOfEntryPoint,
                'BaseOfCode': pe.OPTIONAL_HEADER.BaseOfCode,
                'ImageBase': pe.OPTIONAL_HEADER.ImageBase,
                'SectionAlignment': pe.OPTIONAL_HEADER.SectionAlignment,
                'FileAlignment': pe.OPTIONAL_HEADER.FileAlignment,
                'MajorOperatingSystemVersion': pe.OPTIONAL_HEADER.MajorOperatingSystemVersion,
                'MinorOperatingSystemVersion': pe.OPTIONAL_HEADER.MinorOperatingSystemVersion,
                'MajorImageVersion': pe.OPTIONAL_HEADER.MajorImageVersion,
                'MinorImageVersion': pe.OPTIONAL_HEADER.MinorImageVersion,
                'MajorSubsystemVersion': pe.OPTIONAL_HEADER.MajorSubsystemVersion,
                'MinorSubsystemVersion': pe.OPTIONAL_HEADER.MinorSubsystemVersion,
                'Reserved1': pe.OPTIONAL_HEADER.Reserved1,
                'SizeOfImage': pe.OPTIONAL_HEADER.SizeOfImage,
                'SizeOfHeaders': pe.OPTIONAL_HEADER.SizeOfHeaders,
                'CheckSum': pe.OPTIONAL_HEADER.CheckSum,
                'Subsystem': pe.OPTIONAL_HEADER.Subsystem,
                'DllCharacteristics': pe.OPTIONAL_HEADER.DllCharacteristics,
                'SizeOfStackReserve': pe.OPTIONAL_HEADER.SizeOfStackReserve,
                'SizeOfHeapReserve': pe.OPTIONAL_HEADER.SizeOfHeapReserve,
                'SizeOfHeapCommit': pe.OPTIONAL_HEADER.SizeOfHeapCommit,
                'LoaderFlags': pe.OPTIONAL_HEADER.LoaderFlags,
                'NumberOfRvaAndSizes': pe.OPTIONAL_HEADER.NumberOfRvaAndSizes
            }
            return pd.DataFrame([features])
        except ImportError as e:
            print(f"❌ Missing required EXE libraries: {str(e)}")
            return None
        except Exception as e:
            print(f"❌ Error extracting EXE features: {e}")
            return None

    def calculate_entropy(self, data):
        """Calculate Shannon entropy of data"""
        if not data:
            return 0
        entropy = 0
        data_len = len(data)
        byte_counts = {}
        for byte in data:
            byte_counts[byte] = byte_counts.get(byte, 0) + 1
        for count in byte_counts.values():
            probability = count / data_len
            entropy -= probability * np.log2(probability)
        return entropy

    def count_suspicious_strings(self, data):
        """Count suspicious strings in binary data"""
        suspicious_terms = [
            b'CreateProcess', b'ShellExecute', b'WinExec', b'cmd.exe', b'powershell',
            b'rundll', b'regsvr32', b'http://', b'https://', b'socket', b'connect',
            b'DownloadFile', b'exec', b'eval', b'exploit', b'malware', b'trojan',
            b'backdoor', b'virus', b'payload', b'shellcode', b'encrypt', b'decrypt',
            b'base64', b'inject', b'CreateRemoteThread', b'VirtualAlloc'
        ]
        count = 0
        for term in suspicious_terms:
            count += data.count(term)
        return count

    def extract_multi_exe_features(self, file_path):
        """Extract features from an EXE/DLL for multi-format model"""
        try:
            if self.pefile is None:
                import pefile
                self.pefile = pefile

            pe = self.pefile.PE(file_path)
            features = {}
            for field_name in dir(pe.DOS_HEADER):
                if not field_name.startswith('_'):
                    features[field_name] = getattr(pe.DOS_HEADER, field_name, 0)
            for field_name in dir(pe.FILE_HEADER):
                if not field_name.startswith('_'):
                    features[field_name] = getattr(pe.FILE_HEADER, field_name, 0)
            for field_name in dir(pe.OPTIONAL_HEADER):
                if not field_name.startswith('_'):
                    features[field_name] = getattr(pe.OPTIONAL_HEADER, field_name, 0)
            features['NumberOfSections'] = len(pe.sections)
            data = open(file_path, 'rb').read()
            features['entropy'] = self.calculate_entropy(data)
            features['file_size'] = os.path.getsize(file_path)
            features['contains_executable'] = 1
            features['suspicious_strings'] = self.count_suspicious_strings(data)
            return features
        except ImportError as e:
            print(f"❌ Missing required EXE libraries: {str(e)}")
            return {}
        except Exception as e:
            print(f"❌ Error extracting multi EXE features: {e}")
            return {}

    def extract_doc_features(self, file_path):
        """Extract features from a DOC/DOCX file"""
        try:
            if self.olefile is None:
                import olefile
                self.olefile = olefile
            if file_path.lower().endswith('.docx') and self.docx is None:
                import docx
                self.docx = docx

            features = {field: 0 for field in self.multi_feature_names}
            if self.olefile.isOleFile(file_path):
                ole = self.olefile.OleFileIO(file_path)
                features['contains_macros'] = 1 if 'Macros' in ole.listdir() else 0
                features['object_count'] = len(ole.listdir())
                suspicious_streams = ["VBA", "Macros", "autoexec", "shell", "javascript"]
                count = 0
                for stream in ole.listdir():
                    if any(s.lower() in str(stream).lower() for s in suspicious_streams):
                        count += 1
                features['suspicious_streams'] = count
                try:
                    if ole.exists('\x01Ole10Native'):
                        features['contains_executable'] = 1
                except:
                    features['contains_executable'] = 0
                ole.close()
            elif file_path.lower().endswith('.docx'):
                doc = self.docx.Document(file_path)
                features['object_count'] = len(doc.paragraphs)
                text = ' '.join(p.text for p in doc.paragraphs).lower()
                features['suspicious_strings'] = sum(1 for term in ['javascript', 'http://', 'https://', 'exec'] if term in text)
                features['contains_macros'] = 0
            with open(file_path, 'rb') as f:
                data = f.read()
                features['file_size'] = len(data)
                features['entropy'] = self.calculate_entropy(data)
                features['suspicious_strings'] += self.count_suspicious_strings(data)
            return features
        except ImportError as e:
            print(f"❌ Missing required DOC libraries: {str(e)}")
            return {field: 0 for field in self.multi_feature_names}
        except Exception as e:
            print(f"❌ Error extracting DOC features: {e}")
            return {field: 0 for field in self.multi_feature_names}

    def extract_multi_features(self, file_path, file_type):
        """Extract features based on file type for multi-format model"""
        try:
            features = {field: 0 for field in self.multi_feature_names}
            if file_type == 'exe' or file_path.lower().endswith(('.exe', '.dll')):
                features = self.extract_multi_exe_features(file_path)
                features['file_type'] = 'exe'
            elif any(file_path.lower().endswith(ext) for ext in ['.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx']):
                features = self.extract_doc_features(file_path)
                features['file_type'] = 'doc'
            else:
                with open(file_path, 'rb') as f:
                    data = f.read()
                    features['file_size'] = len(data)
                    features['entropy'] = self.calculate_entropy(data)
                    features['suspicious_strings'] = self.count_suspicious_strings(data)
                features['file_type'] = 'unknown'
            with open(file_path, 'rb') as f:
                sha256 = hashlib.sha256(f.read()).hexdigest()
                features['SHA256'] = sha256
            return features
        except Exception as e:
            print(f"❌ Error extracting multi features: {e}")
            return {field: 0 for field in self.multi_feature_names}

    def select_features(self, features, feature_names):
        """Select top-k features using mutual information"""
        try:
            # Simulate a small dataset for feature selection (single sample is not ideal)
            # Ideally, we'd use a training set, but for single inference, we approximate
            X = np.array([features])
            y = np.array([0])  # Dummy label; mutual_info_classif doesn't use it directly
            selector = SelectKBest(score_func=mutual_info_classif, k=min(self.top_k_features, len(feature_names)))
            selector.fit(X, y)
            selected_indices = selector.get_support(indices=True)
            selected_features = [feature_names[i] for i in selected_indices]
            selected_values = [features[i] for i in selected_indices]
            return selected_features, selected_values, selected_indices
        except Exception as e:
            print(f"❌ Error in feature selection: {e}")
            return feature_names, features, list(range(len(features)))

    def prepare_multi_features(self, data, selected_features=None):
        """Prepare feature vector for multi-format model, using selected features if provided"""
        if selected_features is None:
            selected_features = self.multi_feature_names
        features = []
        for field in selected_features:
            try:
                value = data.get(field, 0)
                features.append(float(value))
            except:
                features.append(0)
        return np.array(features)

    def check_too_many_zeros(self, features):
        """Check if too many features are zero"""
        zero_count = sum(1 for f in features if f == 0)
        zero_ratio = zero_count / len(features)
        return zero_ratio >= self.zero_features_threshold

    def predict(self, file_path):
        """Predict whether a file is malware based on its type"""
        if not os.path.exists(file_path):
            return {"error": f"File not found: {file_path}", "file_path": file_path}

        try:
            # Lazy import python-magic
            if self.magic is None:
                import magic
                self.magic = magic

            file_type = self.magic.from_file(file_path, mime=True)
            ext = os.path.splitext(file_path)[1].lower()
            file_size = os.path.getsize(file_path)
            result = {
                "file_path": file_path,
                "file_type": ext[1:].upper() if ext else "UNKNOWN",
                "file_size": file_size,
                "sha256": hashlib.sha256(open(file_path, 'rb').read()).hexdigest()[:16] + "..."
            }

            if file_type == 'application/pdf' or ext == '.pdf':
                self.load_pdf_model()
                if not self.pdf_model:
                    result["error"] = "PDF model not loaded"
                    return result
                features_df = self.extract_pdf_features(file_path)
                if features_df is None:
                    result["error"] = "Failed to extract PDF features"
                    return result
                prediction = self.pdf_model.predict(features_df)[0]
                confidence = float(self.pdf_model.predict_proba(features_df)[0][prediction]) if hasattr(self.pdf_model, 'predict_proba') else 0.5
                result.update({
                    "is_malware": bool(prediction),
                    "prediction": "MALWARE" if prediction else "BENIGN",
                    "confidence": confidence,
                    "features": features_df.to_dict(orient='records')[0]
                })
                return result

            elif file_type == 'application/x-dosexec' or ext in ('.exe', '.dll'):
                self.load_exe_model()
                if not self.exe_model:
                    result["error"] = "EXE model not loaded"
                    return result
                features_df = self.extract_exe_features(file_path)
                if features_df is None:
                    result["error"] = "Failed to extract EXE features"
                    return result
                prediction = self.exe_model.predict(features_df)[0]
                confidence = float(self.exe_model.predict_proba(features_df)[0][prediction]) if hasattr(self.exe_model, 'predict_proba') else 0.5
                result.update({
                    "is_malware": bool(prediction),
                    "prediction": "MALWARE" if prediction else "BENIGN",
                    "confidence": confidence,
                    "features": features_df.to_dict(orient='records')[0]
                })
                return result

            else:  # Handle DOC, DOCX, XLS, XLSX, PPT, PPTX, and others
                self.load_multi_model()
                if not self.multi_models:
                    result["error"] = "Multi-format model not loaded"
                    return result
                features_dict = self.extract_multi_features(file_path, file_type)
                if not features_dict:
                    result["error"] = "Failed to extract features"
                    return result
                features = self.prepare_multi_features(features_dict)
                
                # Perform feature selection
                selected_feature_names, selected_features, selected_indices = self.select_features(features, self.multi_feature_names)
                is_small_file = file_size < self.small_file_threshold
                too_many_zeros = self.check_too_many_zeros(selected_features)

                if is_small_file or too_many_zeros:
                    model_predictions = {model_name: 0 for model_name in self.multi_models}
                    confidence_scores = {model_name: 0.15 for model_name in self.multi_models}
                    final_prediction = 0
                    overall_confidence = 0.15
                    override_reason = []
                    if is_small_file:
                        override_reason.append(f"File size ({file_size} bytes) below threshold ({self.small_file_threshold} bytes)")
                    if too_many_zeros:
                        zero_count = sum(1 for f in selected_features if f == 0)
                        zero_ratio = zero_count / len(selected_features)
                        override_reason.append(f"Too many zero features in selected set ({zero_ratio:.2%} > {self.zero_features_threshold:.2%})")
                else:
                    # Use selected features for prediction
                    selected_features_vector = self.prepare_multi_features(features_dict, selected_feature_names)
                    model_predictions = {}
                    confidence_scores = {}
                    for model_name, model in self.multi_models.items():
                        model_predictions[model_name] = int(model.predict([selected_features_vector])[0])
                        if hasattr(model, 'predict_proba'):
                            confidence_scores[model_name] = float(model.predict_proba([selected_features_vector])[0][1])
                        else:
                            confidence_scores[model_name] = float(model_predictions[model_name])
                    votes = list(model_predictions.values())
                    final_prediction = 1 if sum(votes) >= len(votes)/2 else 0
                    overall_confidence = sum(confidence_scores.values()) / len(confidence_scores)
                    override_reason = None

                top_features = []
                if hasattr(self.multi_models.get('gb_model', None), 'feature_importances_'):
                    importances = self.multi_models['gb_model'].feature_importances_
                    # Adjust importances to selected features
                    adjusted_importances = [importances[i] for i in selected_indices]
                    indices = np.argsort(adjusted_importances)[::-1]
                    for i in range(min(5, len(indices))):
                        idx = indices[i]
                        if idx < len(selected_feature_names):
                            top_features.append({
                                "name": selected_feature_names[idx],
                                "importance": float(adjusted_importances[idx]),
                                "value": float(selected_features[idx])
                            })

                result.update({
                    "is_malware": bool(final_prediction),
                    "prediction": "MALWARE" if final_prediction else "BENIGN",
                    "confidence": float(overall_confidence),
                    "top_features": top_features,
                    "model_predictions": model_predictions,
                    "model_confidences": confidence_scores,
                    "features": features_dict,
                    "selected_features": selected_feature_names
                })
                if override_reason:
                    result["override_reason"] = override_reason
                return result

        except ImportError as e:
            return {"error": f"Missing required library: {str(e)}", "file_path": file_path}
        except Exception as e:
            return {"error": f"Prediction failed: {str(e)}", "file_path": file_path}

def main():
    parser = argparse.ArgumentParser(description="Unified Malware Detection CLI")
    parser.add_argument("file_path", help="Path to the file to analyze")
    args = parser.parse_args()

    if not os.path.exists(args.file_path):
        print(f"❌ File not found: {args.file_path}")
        sys.exit(1)

    predictor = UnifiedMalwarePredictor()
    result = predictor.predict(args.file_path)

    print("\n🧪 Malware Prediction Result:")
    print(f"📄 File: {result.get('file_path')}")
    print(f"🔖 Type: {result.get('file_type', 'UNKNOWN')}")
    print(f"📏 Size: {result.get('file_size', 0):,} bytes")
    print(f"🔍 SHA256: {result.get('sha256', 'N/A')}")

    if "error" in result:
        print(f"❌ Error: {result['error']}")
    else:
        confidence = result.get('confidence', 0.5)
        if "override_reason" in result:
            print(f"✅ Classified as BENIGN (Automatic override)")
            print(f"⚠️ Override reason:")
            for reason in result["override_reason"]:
                print(f"  - {reason}")
        else:
            if result.get("is_malware"):
                print(f"⚠️ MALWARE DETECTED! ({confidence*100:.2f}% confidence)")
            else:
                print(f"✅ Benign File ({(1-confidence)*100:.2f}% confidence)")

        if "features" in result:
            print("\n🔑 Extracted Features:")
            for key, value in result.get("features", {}).items():
                if key != 'SHA256' and key != 'file_type':
                    print(f"- {key}: {value}")

        if "selected_features" in result:
            print("\n🔍 Selected Features Used:")
            for feature in result["selected_features"]:
                print(f"- {feature}")

        if "top_features" in result and result["top_features"]:
            print("\n🔑 Top Contributing Features:")
            for feature in result["top_features"]:
                print(f"- {feature['name']}: value={feature['value']:.2f}, importance={feature['importance']:.4f}")

        if "model_predictions" in result:
            print("\n🤖 Model Votes:")
            for model, pred in result["model_predictions"].items():
                conf = result["model_confidences"].get(model, 0)
                print(f"- {model}: {'Malware' if pred else 'Benign'} ({conf*100:.2f}% confidence)")

if __name__ == "__main__":
    main()