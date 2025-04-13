import sys
import os
import pickle
import numpy as np
import pandas as pd
import pefile
import olefile
import magic
import hashlib
from sklearn.ensemble import GradientBoostingClassifier, RandomForestClassifier

# === Step 1: Define the EnhancedMalwarePredictor class ===
class EnhancedMalwarePredictor:
    """Class to predict malware using trained models from EnhancedMalwareDetector"""
    
    def __init__(self, model_path="models/malware_models.pkl", small_file_threshold=10240, zero_features_threshold=0.7):
        """Initialize with path to trained models"""
        self.models = None
        self.feature_names = None
        self.model_info = None
        self.small_file_threshold = small_file_threshold  # Files smaller than 10KB are considered small
        self.zero_features_threshold = zero_features_threshold  # If 70% or more features are 0, consider benign
        self.load_models(model_path)
    
    def load_models(self, model_path):
        """Load trained malware detection models"""
        try:
            with open(model_path, "rb") as f:
                self.models = pickle.load(f)
            
            # Load model info if available
            try:
                with open(model_path.replace(".pkl", "_info.json"), "r") as f:
                    self.model_info = json.load(f)
            except:
                self.model_info = {"metrics": {"accuracy": "unknown"}}
            
            # Define feature names (same as in EnhancedMalwareDetector)
            self.feature_names = [
                'e_magic', 'e_cblp', 'e_cp', 'e_crlc', 'e_cparhdr', 'e_minalloc', 'e_maxalloc',
                'e_ss', 'e_sp', 'e_csum', 'e_ip', 'e_cs', 'e_lfarlc', 'e_ovno', 'e_oemid',
                'e_oeminfo', 'e_lfanew', 'Machine', 'NumberOfSections', 'TimeDateStamp',
                'PointerToSymbolTable', 'NumberOfSymbols', 'SizeOfOptionalHeader', 'Characteristics',
                'Magic', 'MajorLinkerVersion', 'MinorLinkerVersion', 'SizeOfCode',
                'SizeOfInitializedData', 'SizeOfUninitializedData', 'AddressOfEntryPoint',
                'BaseOfCode', 'ImageBase', 'SectionAlignment', 'FileAlignment',
                'MajorOperatingSystemVersion', 'MinorOperatingSystemVersion', 'MajorImageVersion',
                'MinorImageVersion', 'MajorSubsystemVersion', 'MinorSubsystemVersion', 'Reserved1',
                'SizeOfImage', 'SizeOfHeaders', 'CheckSum', 'Subsystem', 'DllCharacteristics',
                'SizeOfStackReserve', 'SizeOfHeapReserve', 'SizeOfHeapCommit', 'LoaderFlags',
                'NumberOfRvaAndSizes', 'file_size', 'entropy', 'contains_macros', 'suspicious_streams',
                'object_count', 'javascript_count', 'urls_count', 'suspicious_strings',
                'encryption_indicator', 'obfuscation_level', 'contains_executable', 'unusual_metadata'
            ]
            
            print(f"✅ Models loaded successfully from {model_path}")
            return True
        
        except Exception as e:
            print(f"❌ Error loading models: {str(e)}")
            return False
    
    def calculate_entropy(self, data):
        """Calculate Shannon entropy of data"""
        if not data:
            return 0
        entropy = 0
        data_len = len(data)
        byte_counts = {}
        for byte in data:
            byte_counts[byte] = byte_counts.get(byte, 0) + 1
        for count in byte_counts.values():
            probability = count / data_len
            entropy -= probability * np.log2(probability)
        return entropy
    
    def count_suspicious_strings(self, data):
        """Count suspicious strings in binary data"""
        suspicious_terms = [
            b'CreateProcess', b'ShellExecute', b'WinExec', b'cmd.exe', b'powershell',
            b'rundll', b'regsvr32', b'http://', b'https://', b'socket', b'connect',
            b'DownloadFile', b'exec', b'eval', b'exploit', b'malware', b'trojan',
            b'backdoor', b'virus', b'payload', b'shellcode', b'encrypt', b'decrypt',
            b'base64', b'inject', b'CreateRemoteThread', b'VirtualAlloc'
        ]
        count = 0
        for term in suspicious_terms:
            count += data.count(term)
        return count
    
    def extract_features_from_exe(self, file_path):
        """Extract features from an EXE/PE file"""
        try:
            pe = pefile.PE(file_path)
            features = {}
            
            # Extract DOS Header features
            for field_name in dir(pe.DOS_HEADER):
                if not field_name.startswith('_'):
                    try:
                        features[field_name] = getattr(pe.DOS_HEADER, field_name)
                    except:
                        features[field_name] = 0
            
            # Extract File Header features
            for field_name in dir(pe.FILE_HEADER):
                if not field_name.startswith('_'):
                    try:
                        features[field_name] = getattr(pe.FILE_HEADER, field_name)
                    except:
                        features[field_name] = 0
            
            # Extract Optional Header features
            for field_name in dir(pe.OPTIONAL_HEADER):
                if not field_name.startswith('_'):
                    try:
                        features[field_name] = getattr(pe.OPTIONAL_HEADER, field_name)
                    except:
                        features[field_name] = 0
            
            # Add section information
            features['NumberOfSections'] = len(pe.sections)
            
            # Calculate file entropy
            data = open(file_path, 'rb').read()
            features['entropy'] = self.calculate_entropy(data)
            
            # Add file size
            features['file_size'] = os.path.getsize(file_path)
            
            # Other PE specific features
            features['contains_executable'] = 1
            features['suspicious_strings'] = self.count_suspicious_strings(data)
            
            return features
        
        except Exception as e:
            print(f"❌ Error extracting features from EXE: {str(e)}")
            return {}
    
    def extract_features_from_doc(self, file_path):
        """Extract features from a DOC/DOCX file"""
        try:
            features = {field: 0 for field in self.feature_names}

            # Handle legacy .doc (OLE format)
            if olefile.isOleFile(file_path):
                ole = olefile.OleFileIO(file_path)

                features['contains_macros'] = 1 if 'Macros' in ole.listdir() else 0
                features['object_count'] = len(ole.listdir())

                suspicious_streams = ["VBA", "Macros", "autoexec", "shell", "javascript"]
                count = 0
                for stream in ole.listdir():
                    if any(s.lower() in str(stream).lower() for s in suspicious_streams):
                        count += 1
                features['suspicious_streams'] = count

                try:
                    if ole.exists('\x01Ole10Native'):
                        features['contains_executable'] = 1
                except:
                    features['contains_executable'] = 0
                ole.close()

            # Handle .docx (OpenXML format)
            elif file_path.lower().endswith('.docx'):
                try:
                    from docx import Document
                    doc = Document(file_path)
                    # Basic .docx analysis (expand as needed)
                    features['object_count'] = len(doc.paragraphs)
                    # Check for embedded scripts or links (simplified)
                    text = ' '.join(p.text for p in doc.paragraphs).lower()
                    features['suspicious_strings'] = sum(1 for term in ['javascript', 'http://', 'https://', 'exec'] if term in text)
                    features['contains_macros'] = 0  # .docx typically doesn't embed VBA macros
                except Exception as e:
                    print(f"⚠️ Warning: Could not parse .docx: {str(e)}")

            # Common features
            with open(file_path, 'rb') as f:
                data = f.read()
                features['file_size'] = len(data)
                features['entropy'] = self.calculate_entropy(data)
                features['suspicious_strings'] += self.count_suspicious_strings(data)

            return features
    
        except Exception as e:
            print(f"❌ Error extracting features from DOC: {str(e)}")
            return features
    
    def extract_features_from_file(self, file_path):
        """Extract features based on file type"""
        try:
            # Detect file type
            file_type = magic.from_file(file_path, mime=True)
            
            # Extract features based on file type
            if file_type == 'application/x-dosexec' or file_path.lower().endswith(('.exe', '.dll')):
                features = self.extract_features_from_exe(file_path)
                features['file_type'] = 'exe'
            
            elif any(file_path.lower().endswith(ext) for ext in ['.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx']):
                features = self.extract_features_from_doc(file_path)
                features['file_type'] = 'doc'
            
            else:
                # Basic metrics for unknown files
                features = {field: 0 for field in self.feature_names}
                with open(file_path, 'rb') as f:
                    data = f.read()
                    features['file_size'] = len(data)
                    features['entropy'] = self.calculate_entropy(data)
                    features['suspicious_strings'] = self.count_suspicious_strings(data)
                features['file_type'] = 'unknown'
            
            # Compute SHA256 hash
            with open(file_path, 'rb') as f:
                sha256 = hashlib.sha256(f.read()).hexdigest()
                features['SHA256'] = sha256
            
            return features
        
        except Exception as e:
            print(f"❌ Error extracting features from file: {str(e)}")
            return {}
    
    def prepare_features(self, data):
        """Prepare feature vector from file data"""
        features = []
        for field in self.feature_names:
            try:
                value = data.get(field, 0)
                try:
                    features.append(float(value))
                except:
                    features.append(0)
            except:
                features.append(0)
        return np.array(features)
    
    def check_too_many_zeros(self, features):
        """Check if too many features are zero (indicating likely benign file)"""
        zero_count = sum(1 for f in features if f == 0)
        zero_ratio = zero_count / len(features)
        return zero_ratio >= self.zero_features_threshold
    
    def predict(self, file_path):
        """Predict whether a file is malware"""
        if not self.models:
            return {"error": "No models loaded", "prediction": None}
        
        try:
            # Extract features
            features_dict = self.extract_features_from_file(file_path)
            if not features_dict:
                return {"error": "Failed to extract features", "file_path": file_path}
            
            # Get file size
            file_size = features_dict.get('file_size', 0)
            
            # Prepare features for model
            features = self.prepare_features(features_dict)
            
            # Check if file is too small or if too many features are zero
            is_small_file = file_size < self.small_file_threshold
            too_many_zeros = self.check_too_many_zeros(features)
            
            # If file is too small or has too many zero features, classify as benign
            if is_small_file or too_many_zeros:
                # Force benign prediction for small files or files with too many zeros
                model_predictions = {model_name: 0 for model_name in self.models}
                confidence_scores = {model_name: 0.15 for model_name in self.models}  # Low confidence
                final_prediction = 0
                overall_confidence = 0.15
                
                # Include the reason for the benign classification
                override_reason = []
                if is_small_file:
                    override_reason.append(f"File size ({file_size} bytes) below threshold ({self.small_file_threshold} bytes)")
                if too_many_zeros:
                    zero_count = sum(1 for f in features if f == 0)
                    zero_ratio = zero_count / len(features)
                    override_reason.append(f"Too many zero features ({zero_ratio:.2%} > {self.zero_features_threshold:.2%})")
            else:
                # Normal prediction flow for regular files
                # Get predictions from each model
                model_predictions = {}
                confidence_scores = {}
                for model_name, model in self.models.items():
                    model_predictions[model_name] = int(model.predict([features])[0])
                    if hasattr(model, 'predict_proba'):
                        confidence_scores[model_name] = float(model.predict_proba([features])[0][1])
                    else:
                        confidence_scores[model_name] = float(model_predictions[model_name])
                
                # Majority voting
                votes = list(model_predictions.values())
                final_prediction = 1 if sum(votes) >= len(votes)/2 else 0
                overall_confidence = sum(confidence_scores.values()) / len(confidence_scores)
                override_reason = None
            
            # Feature importance analysis
            top_features = []
            if hasattr(self.models.get('gb_model', None), 'feature_importances_'):
                importances = self.models['gb_model'].feature_importances_
                indices = np.argsort(importances)[::-1]
                for i in range(min(5, len(indices))):
                    idx = indices[i]
                    if idx < len(self.feature_names):
                        top_features.append({
                            "name": self.feature_names[idx],
                            "importance": float(importances[idx]),
                            "value": float(features[idx])
                        })
            
            # Prepare result
            result = {
                "file_path": file_path,
                "prediction": int(final_prediction),
                "is_malware": bool(final_prediction),
                "confidence": float(overall_confidence),
                "file_type": features_dict.get('file_type', 'unknown'),
                "sha256": features_dict.get('SHA256', ''),
                "file_size": file_size,
                "top_features": top_features,
                "model_predictions": model_predictions,
                "model_confidences": confidence_scores
            }
            
            # Add override reason if applicable
            if override_reason:
                result["override_reason"] = override_reason
            
            return result
        
        except Exception as e:
            return {"error": str(e), "file_path": file_path, "prediction": None}

# === Step 2: CLI input and prediction ===
if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python predict_malware.py <path_to_file>")
        print("Supported files: .exe, .dll, .doc, .docx, .xls, .xlsx, .ppt, .pptx")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        sys.exit(1)
    
    # Initialize predictor
    predictor = EnhancedMalwarePredictor()
    
    # Make prediction
    result = predictor.predict(file_path)
    
    # Display results
    print("\n🧪 Malware Prediction Result:")
    print(f"📄 File: {result.get('file_path')}")
    print(f"🔖 Type: {result.get('file_type', 'unknown').upper()}")
    print(f"📏 Size: {result.get('file_size', 0):,} bytes")
    print(f"🔍 SHA256: {result.get('sha256', 'N/A')[:16]}... (truncated)")
    
    if "error" in result:
        print(f"❌ Error: {result['error']}")
    else:
        # Check for override reasons
        if "override_reason" in result:
            print(f"✅ Classified as BENIGN (Automatic override)")
            print(f"⚠️ Override reason:")
            for reason in result["override_reason"]:
                print(f"  - {reason}")
        else:
            if result["is_malware"]:
                print(f"⚠️ MALWARE DETECTED! ({result['confidence']*100:.2f}% confidence)")
            else:
                print(f"✅ Benign File ({(1-result['confidence'])*100:.2f}% confidence)")
        
        print("\n🔑 Top Contributing Features:")
        for feature in result.get("top_features", []):
            print(f"- {feature['name']}: value={feature['value']:.2f}, importance={feature['importance']:.4f}")
        
        print("\n🤖 Model Votes:")
        for model, pred in result.get("model_predictions", {}).items():
            conf = result.get("model_confidences", {}).get(model, 0)
            print(f"- {model}: {'Malware' if pred else 'Benign'} ({conf*100:.2f}% confidence)")